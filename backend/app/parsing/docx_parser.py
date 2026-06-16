from docx import Document


def parse_docx(file_path: str) -> str:
    """Extract text from a Word document. Returns empty string on failure."""
    try:
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = []
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    tables.append(row_text)
        return '\n'.join(paragraphs + tables).strip()
    except Exception as e:
        return f'[DOCX parse error: {e}]'
